"""Generate the daily manual-Westlaw download worklist.

The 1953-~1965 slice of the gap era (N.W.2d vols below 139) is NOT in
the court-sourced NW-cite archive, so the only path to a high-authority
text is a manual Westlaw pull. Completion of the gap era requires
acquiring every such opinion from Westlaw; that fetch is manual and
capped (~500/day). This builds the day's list as a Word doc split into
batches of <=100 citations the user can hand to Westlaw, leading with
the opinions that bear the strongest indications of error.

Target pool: gap_1953_1996 opinions still in state 'unvalidated' or
'flagged' (the court-archive promotion already drained the rest to
'corrected'/'cross_checked'), that carry an N.W. citation and are not
already on a prior worklist.

Priority (quality_score is uninformative for the NW2d OCR era per
TODO §2, so it is NOT used):
  1. flagged before unvalidated — a flag is a known court-archive
     divergence/ambiguity, the strongest error signal.
  2. missing/blank author before populated — metadata anomaly.
  3. higher inbound-citation count first — get the load-bearing
     precedent right before the obscure.
  4. older first — chip the era away chronologically.

State: every listed opinion is recorded in westlaw_requests so the
next day's run excludes it. Receiving/ingesting the returned .docs is
a separate future step (received_at/received_path columns reserved).

Usage:
  python -m ndcourts_mcp.westlaw_worklist --status
  python -m ndcourts_mcp.westlaw_worklist                 # preview 500
  python -m ndcourts_mcp.westlaw_worklist --apply          # record + .docx
  python -m ndcourts_mcp.westlaw_worklist --limit 300 --apply
"""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document

from .db import DEFAULT_DB_PATH, get_connection, log_provenance

DEFAULT_LIMIT = 500          # opinions per day (manual Westlaw cap)
DEFAULT_PER_SECTION = 100    # citations per Westlaw batch
WORKLIST_DIR = Path("worklists")

# An unreceived listed row is "in flight" for this many days (the manual
# Westlaw download/receive cycle) before it ages back into the pool.
# Without this gate the pool re-lists every recently-listed row on every
# --apply (a re-list storm); with it, rows are surfaced again only once
# they are genuinely stale.
RELIST_AFTER_DAYS = 7

# An opinion still needs a Westlaw pull when it has NOT been received AND
# it is not currently in flight (listed within RELIST_AFTER_DAYS, not yet
# back) AND it is not one half of a §6 duplicate-row pair (same N.W. cite
# + same date_filed as another opinion — one Find&Print doc matches both
# rows, so re-listing wastes a pull until §6 collapses the pair).
# Replaces the old blanket `id NOT IN westlaw_requests`, which excluded
# every listed opinion forever regardless of whether it ever came back.
# RELIST_AFTER_DAYS is an int constant we control — safe to inline.
_OUTSTANDING_PREDICATE = f"""
      AND o.id NOT IN (SELECT opinion_id FROM westlaw_requests
                        WHERE received_at IS NOT NULL)
      AND o.id NOT IN (SELECT opinion_id FROM westlaw_requests
                        WHERE received_at IS NULL
                          AND listed_at >=
                              datetime('now','-{RELIST_AFTER_DAYS} days'))
      AND NOT EXISTS (
            SELECT 1 FROM citations c2
              JOIN citations c3 ON c3.citation = c2.citation
                               AND c3.opinion_id <> c2.opinion_id
              JOIN opinions o3  ON o3.id = c3.opinion_id
                               AND o3.date_filed = o.date_filed
             WHERE c2.opinion_id = o.id
               AND c2.citation LIKE '% N.W.%')
"""

POOL_SQL = """
    SELECT o.id, o.case_name, o.date_filed, v.crosscheck_state,
           (o.author IS NULL OR TRIM(o.author) = '') AS missing_author,
           (SELECT c.citation FROM citations c
             WHERE c.opinion_id = o.id AND c.citation LIKE '% N.W.%'
             ORDER BY c.citation LIMIT 1) AS nw_cite,
           (SELECT COUNT(*) FROM cited_by cb
             WHERE cb.cited_opinion_id = o.id) AS inbound
    FROM validation_status v
    JOIN opinions o ON o.id = v.opinion_id
    WHERE v.era_tier = 'gap_1953_1996'
      AND v.crosscheck_state IN ('unvalidated', 'flagged')
""" + _OUTSTANDING_PREDICATE + """
      AND EXISTS (SELECT 1 FROM citations c
                   WHERE c.opinion_id = o.id AND c.citation LIKE '% N.W.%')
    ORDER BY
      CASE v.crosscheck_state WHEN 'flagged' THEN 0 ELSE 1 END,
      missing_author DESC,
      inbound DESC,
      o.date_filed ASC
"""


def _ensure_table(conn) -> None:
    conn.execute("""
        CREATE TABLE IF NOT EXISTS westlaw_requests (
            opinion_id    INTEGER PRIMARY KEY REFERENCES opinions(id),
            nw_cite       TEXT NOT NULL,
            doc_batch     TEXT NOT NULL,
            listed_at     TEXT NOT NULL DEFAULT
                          (strftime('%Y-%m-%dT%H:%M:%S','now')),
            received_at   TEXT,
            received_path TEXT
        )
    """)


def _reconcile_received(conn) -> int:
    """Backfill received_at for listed rows whose opinion now carries
    Westlaw bound text (filled via a shared-cite page-mate or hand-
    adjudication under a different oid than the one listed). Without
    this they look 'unreceived' forever and would be wrongly re-listed.
    Conservative: only touches rows where the opinion is already
    westlaw-primary, i.e. already validated. Idempotent."""
    cur = conn.execute("""
        UPDATE westlaw_requests
           SET received_at = strftime('%Y-%m-%dT%H:%M:%S','now'),
               received_path = COALESCE(received_path,
                 (SELECT o.source_path FROM opinions o
                   WHERE o.id = westlaw_requests.opinion_id))
         WHERE received_at IS NULL
           AND opinion_id IN (SELECT id FROM opinions
                               WHERE source_reporter = 'westlaw')
    """)
    return cur.rowcount


def _pool(conn, limit: int) -> list:
    return conn.execute(POOL_SQL + " LIMIT ?", (limit,)).fetchall()


def _write_docx(rows: list, out_path: Path, per_section: int,
                batch_date: str) -> None:
    doc = Document()
    doc.add_heading("Westlaw download worklist", level=0)
    n_sec = (len(rows) + per_section - 1) // per_section
    doc.add_paragraph(
        f"Batch date: {batch_date}.  {len(rows)} opinions in "
        f"{n_sec} section(s) of <= {per_section}.  "
        f"Look each up in Westlaw by the N.W. citation; the case name "
        f"and filing date are for your cross-reference. Ordered "
        f"highest-error-indication first."
    )
    for i in range(0, len(rows), per_section):
        chunk = rows[i:i + per_section]
        sec = i // per_section + 1
        doc.add_heading(
            f"Section {sec} of {n_sec} — {len(chunk)} citations", level=1
        )
        for r in chunk:
            flag = "  [flagged]" if r["crosscheck_state"] == "flagged" else ""
            doc.add_paragraph(
                f'{r["nw_cite"]}  —  {r["case_name"]} '
                f'({r["date_filed"]}){flag}',
                style="List Number",
            )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def status(conn) -> None:
    _ensure_table(conn)
    reconciled = _reconcile_received(conn)
    conn.commit()
    pool_total = conn.execute(
        "SELECT COUNT(*) FROM validation_status v JOIN opinions o "
        "ON o.id=v.opinion_id WHERE v.era_tier='gap_1953_1996' "
        "AND v.crosscheck_state IN ('unvalidated','flagged') "
        "AND EXISTS (SELECT 1 FROM citations c WHERE c.opinion_id=o.id "
        "AND c.citation LIKE '% N.W.%')"
    ).fetchone()[0]
    listed = conn.execute("SELECT COUNT(*) FROM westlaw_requests").fetchone()[0]
    received = conn.execute(
        "SELECT COUNT(*) FROM westlaw_requests WHERE received_at IS NOT NULL"
    ).fetchone()[0]
    # §6 dup-pair-blocked: listed, unreceived, not westlaw — held until §6.
    dup_blocked = conn.execute(
        "SELECT COUNT(*) FROM westlaw_requests w JOIN opinions o "
        "ON o.id=w.opinion_id WHERE w.received_at IS NULL "
        "AND o.source_reporter<>'westlaw' AND EXISTS ("
        "SELECT 1 FROM citations c2 JOIN citations c3 "
        "ON c3.citation=c2.citation AND c3.opinion_id<>c2.opinion_id "
        "JOIN opinions o3 ON o3.id=c3.opinion_id "
        "AND o3.date_filed=o.date_filed "
        "WHERE c2.opinion_id=o.id AND c2.citation LIKE '% N.W.%')"
    ).fetchone()[0]
    remaining = conn.execute(
        "SELECT COUNT(*) FROM validation_status v JOIN opinions o "
        "ON o.id=v.opinion_id WHERE v.era_tier='gap_1953_1996' "
        "AND v.crosscheck_state IN ('unvalidated','flagged') "
        + _OUTSTANDING_PREDICATE +
        " AND EXISTS (SELECT 1 FROM citations c WHERE c.opinion_id=o.id "
        "AND c.citation LIKE '% N.W.%')"
    ).fetchone()[0]
    print("=== Westlaw worklist status ===")
    print(f"  gap-era pool (needs Westlaw):  {pool_total}")
    print(f"  listed on prior worklists:     {listed}")
    print(f"  received back:                 {received}"
          f"  (+{reconciled} reconciled this run)")
    print(f"  §6 dup-pair blocked (held):    {dup_blocked}")
    print(f"  remaining to list:             {remaining}")


def export_citations(conn, doc_batch: str | None, per_block: int,
                     out_path: Path) -> tuple[int, int, Path]:
    """Write a paste-ready citations-only .txt from the RECORDED listing
    in westlaw_requests (source of truth — stays in sync with what was
    listed; never re-selects the pool). Per Westlaw: no case names,
    citations semicolon-separated, batched at the per-request cap."""
    _ensure_table(conn)
    if doc_batch is None:
        row = conn.execute(
            "SELECT doc_batch FROM westlaw_requests "
            "ORDER BY listed_at DESC LIMIT 1").fetchone()
        if not row:
            raise SystemExit("westlaw_requests is empty — generate a "
                             "worklist with --apply first.")
        doc_batch = row[0]
    cites = [r[0] for r in conn.execute(
        "SELECT nw_cite FROM westlaw_requests WHERE doc_batch=? "
        "ORDER BY listed_at, opinion_id", (doc_batch,))]
    n_blocks = (len(cites) + per_block - 1) // per_block
    lines = [
        f"Westlaw Find & Print — batch {doc_batch}",
        f"{len(cites)} citations in {n_blocks} block(s) of <= {per_block}.",
        "Paste ONE block per Find & Print request (citations only, "
        "semicolon-separated, no case names — per Westlaw's instruction).",
        "",
    ]
    for i in range(0, len(cites), per_block):
        block = cites[i:i + per_block]
        lines.append(f"===== Block {i // per_block + 1} of {n_blocks} "
                     f"({len(block)} citations) =====")
        lines.append("; ".join(block))
        lines.append("")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return len(cites), n_blocks, out_path


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--db", type=Path, default=DEFAULT_DB_PATH)
    ap.add_argument("--limit", type=int, default=DEFAULT_LIMIT)
    ap.add_argument("--per-section", type=int, default=DEFAULT_PER_SECTION)
    ap.add_argument("--out", type=Path)
    ap.add_argument("--apply", action="store_true",
                    help="Record opinions as listed + write the .docx "
                         "(default: preview only, records nothing)")
    ap.add_argument("--status", action="store_true")
    ap.add_argument("--export-citations", action="store_true",
                    help="Write a paste-ready semicolon-separated "
                         "citations-only .txt from the recorded listing")
    ap.add_argument("--batch", help="doc_batch to export "
                    "(default: most recent)")
    args = ap.parse_args()

    conn = get_connection(args.db)
    try:
        if args.status:
            status(conn)
            return
        if args.export_citations:
            today = date.today().isoformat()
            out = args.out or (WORKLIST_DIR
                               / f"westlaw-{args.batch or today}-cites.txt")
            n, nb, p = export_citations(
                conn, args.batch, args.per_section, out)
            print(f"Wrote {n} citations in {nb} block(s) of "
                  f"<= {args.per_section} -> {p}")
            return
        _ensure_table(conn)
        reconciled = _reconcile_received(conn)
        if reconciled:
            print(f"Reconciled {reconciled} listed row(s) now westlaw-primary "
                  f"(filled via page-mate/hand-adj) -> marked received.")
        rows = _pool(conn, args.limit)
        if not rows:
            print("Pool empty — every gap-era opinion is listed or drained.")
            return
        today = date.today().isoformat()
        n_sec = (len(rows) + args.per_section - 1) // args.per_section
        out = args.out or (WORKLIST_DIR / f"westlaw-{today}.docx")

        if not args.apply:
            print(f"DRY RUN — would list {len(rows)} opinions in {n_sec} "
                  f"section(s), write {out}. Records nothing.")
            print("  top 5:")
            for r in rows[:5]:
                print(f"    {r['nw_cite']:<18} {r['case_name'][:46]} "
                      f"[{r['crosscheck_state']}] inbound={r['inbound']}")
            print("  Pass --apply to record + generate the Word doc.")
            return

        _write_docx(rows, out, args.per_section, today)
        conn.executemany(
            "INSERT INTO westlaw_requests (opinion_id, nw_cite, doc_batch) "
            "VALUES (?, ?, ?) "
            "ON CONFLICT(opinion_id) DO UPDATE SET "
            "  doc_batch=excluded.doc_batch, nw_cite=excluded.nw_cite, "
            "  listed_at=strftime('%Y-%m-%dT%H:%M:%S','now'), "
            "  received_at=NULL, received_path=NULL",
            [(r["id"], r["nw_cite"], today) for r in rows],
        )
        log_provenance(
            conn,
            operation="westlaw_worklist",
            command=f"python -m ndcourts_mcp.westlaw_worklist --apply "
                    f"--limit {args.limit}",
            rows_affected=len(rows),
            notes=(f"listed {len(rows)} gap-era opinions for manual Westlaw "
                   f"pull, batch {today}, {n_sec} sections of "
                   f"<= {args.per_section}; doc {out}"),
        )
        conn.commit()
        print(f"Listed {len(rows)} opinions (batch {today}); wrote {out}")
    finally:
        conn.close()


if __name__ == "__main__":
    main()
